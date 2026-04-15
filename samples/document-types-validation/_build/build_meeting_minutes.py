"""Operating Committee Agenda + Minutes — combined DOCX."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from nextdecade_brand import (
    new_document, add_heading, add_body, add_bullet, add_runtext, add_table,
    set_footer, set_header, add_horizontal_rule,
    NAVY, BLACK, ORANGE, GREY,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build(output: str):
    doc = new_document()
    set_header(doc, "NextDecade Corporation",
               "Operating Committee  |  April 10, 2026")

    title = doc.add_paragraph()
    add_runtext(title, "OPERATING COMMITTEE", bold=True, size=20, color=NAVY)
    sub = doc.add_paragraph()
    add_runtext(sub, "Agenda and Minutes  -  April 10, 2026",
                italic=True, size=12, color=ORANGE)
    add_horizontal_rule(doc, color=NAVY)

    # Header block
    t = doc.add_table(rows=4, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(1.4)
    t.columns[1].width = Inches(5.1)
    rows = [
        ("DATE/TIME:", "April 10, 2026  -  9:00 a.m. to 10:30 a.m. Central"),
        ("LOCATION:", "1000 Louisiana Street, Suite 3900, Houston - Room 3901 and Teams"),
        ("CHAIR:", "Matt Schatzman, Chairman and Chief Executive Officer"),
        ("SCRIBE:", "L. Reynolds, Executive Assistant"),
    ]
    for i, (label, value) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.text = ""
        add_runtext(c0.paragraphs[0], label, bold=True, size=10, color=NAVY)
        c1.text = ""
        add_runtext(c1.paragraphs[0], value, size=10, color=BLACK)

    # Attendees
    add_heading(doc, "Attendees", level=2)
    add_body(doc, "Present: M. Schatzman (Chair), B. Smith (COO), K. Nguyen (CFO), "
                  "S. Patel (VP Supply Chain), R. Chen (CIO), J. Doe (General Counsel), "
                  "A. Gomez (VP HSSE), D. Park (VP Operations).")
    add_body(doc, "Absent: T. Williams (VP Commercial) - excused.")

    # Agenda
    add_heading(doc, "Agenda", level=2)
    for i, item in enumerate([
        "Safety Moment and HSSE report (A. Gomez, 10 min)",
        "Construction progress - Q1 2026 (B. Smith, 20 min)",
        "Q1 financial summary and liquidity (K. Nguyen, 15 min)",
        "Commercial update - Phase 2 SPA negotiations (T. Williams written report, 10 min)",
        "IT and cyber readiness (R. Chen, 10 min)",
        "Legal and regulatory update (J. Doe, 15 min)",
        "Open items and close (M. Schatzman, 10 min)",
    ], 1):
        add_bullet(doc, f"{i}. {item}")

    # Minutes
    add_heading(doc, "Minutes", level=2)

    add_heading(doc, "1. Safety Moment and HSSE Report", level=3)
    add_body(doc,
        "A. Gomez opened with a safety moment on hot work and fire watch "
        "stewardship. The near-miss event HW-2026-0143 (April 14, 2026) "
        "was reviewed; all corrective actions are complete. The Total "
        "Recordable Incident Rate for Q1 2026 is 0.42, below the industry "
        "benchmark of 0.65.")

    add_heading(doc, "2. Construction Progress", level=3)
    add_body(doc,
        "B. Smith reported that Phase 1 construction is 68 percent complete. "
        "Engineering is 99 percent, procurement is 88 percent. All long-lead "
        "equipment is on site or in transit. Commissioning of non-cryogenic "
        "utilities starts July 2026. First LNG remains on schedule for "
        "the second half of 2027.")

    add_heading(doc, "3. Q1 Financial Summary", level=3)
    add_body(doc,
        "K. Nguyen reported that the Company remains on budget at $18.4B "
        "for Phase 1. Quarter-end liquidity is adequate through first LNG. "
        "The Committee discussed the Phase 2 financing roadmap.")

    add_heading(doc, "4. Commercial Update", level=3)
    add_body(doc,
        "T. Williams's written report was distributed in advance. Three "
        "investment-grade counterparties are in advanced SPA negotiations "
        "for Phase 2 volumes. Committee requested a verbal update at the "
        "next meeting.")

    add_heading(doc, "5. IT and Cyber Readiness", level=3)
    add_body(doc,
        "R. Chen reported completion of the annual tabletop exercise. No "
        "material exposure identified. Endpoint configuration baseline is "
        "current. VPN and remote access controls validated; see related "
        "Remote Work Guidance issued April 15, 2026.")

    add_heading(doc, "6. Legal and Regulatory Update", level=3)
    add_body(doc,
        "J. Doe confirmed that the FERC Phase 2 authorization remains in "
        "effect. The DOE non-FTA confirmation is expected in Q2 2026. No "
        "new material litigation. The Records Retention Standard was "
        "approved for issuance effective April 15, 2026.")

    # Decisions and actions
    add_heading(doc, "Decisions Made", level=2)
    add_bullet(doc, "Records Retention Standard - approved for issuance.")
    add_bullet(doc, "Remote Work Guidance - acknowledged.")
    add_bullet(doc, "Commercial verbal update - added to May 2026 agenda.")

    add_heading(doc, "Action Items", level=2)
    add_table(
        doc,
        headers=["#", "Action", "Owner", "Due"],
        rows=[
            ["1", "Circulate finalized Records Retention Standard to all department heads.", "J. Doe", "Apr 22, 2026"],
            ["2", "Commercial Phase 2 SPA update at May OpsComm.", "T. Williams", "May 8, 2026"],
            ["3", "Tabletop exercise after-action report to OpsComm.", "R. Chen", "May 8, 2026"],
            ["4", "HW-2026-0143 corrective action closeout brief.", "A. Gomez", "May 1, 2026"],
        ],
    )

    # Adjournment
    add_heading(doc, "Adjournment", level=2)
    add_body(doc, "The meeting adjourned at 10:28 a.m. Central. The next Operating "
                  "Committee meeting is scheduled for May 8, 2026.")

    # Signature
    doc.add_paragraph()
    sig = doc.add_paragraph()
    add_runtext(sig, "Minutes submitted by:  L. Reynolds, Executive Assistant",
                size=10, color=GREY)
    sig2 = doc.add_paragraph()
    add_runtext(sig2, "Approved by:  Matt Schatzman, Chair", size=10, color=GREY)

    set_footer(doc, classification=True, mission=True, page_numbers=True)
    doc.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    build(sys.argv[1])
