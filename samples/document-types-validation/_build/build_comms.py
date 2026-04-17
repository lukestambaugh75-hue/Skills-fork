"""Internal comms artifacts — All-Hands Email + NEXT Digest Newsletter.

Voice: casual/internal per Writing Style Guide §5a. Contractions used ("we're", "it's").
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from nextdecade_brand import (
    new_document, add_heading, add_body, add_bullet, add_runtext,
    set_footer, add_horizontal_rule,
    NAVY, BLACK, ORANGE, GREEN, GREY,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_all_hands_email(output: Path):
    doc = new_document()

    # Department banner (orange bar)
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(banner, "OFFICE OF THE CEO", bold=True, size=10, color=ORANGE)

    # Email headers block
    def row(label, value):
        p = doc.add_paragraph()
        add_runtext(p, label, bold=True, size=10, color=NAVY)
        add_runtext(p, "  " + value, size=10, color=BLACK)

    row("From:", "Matt Schatzman <matt.schatzman@next-decade.com>")
    row("To:", "All NextDecade <allnextdecade@next-decade.com>")
    row("Sent:", "April 15, 2026  3:30 p.m. Central")
    row("Subject:", "Q1 check-in - 68%, safety, and what's next")

    add_horizontal_rule(doc, color=NAVY)

    # Body — casual/internal voice with contractions
    add_body(doc, "Team,")
    add_body(doc,
        "Quick update as we close out Q1 2026. Rio Grande LNG Phase 1 is "
        "now 68% complete. That's a real milestone, and it's because of "
        "the work every one of you is doing. Thank you.")

    add_heading(doc, "Where we are on construction", level=3)
    add_body(doc,
        "Engineering is basically done (99%). Procurement is almost there "
        "(88%). Construction is 52% across all three trains. The Train 1 "
        "cold box went up this quarter without incident - nothing routine "
        "about that, and it's a huge win for the commissioning team.")
    add_body(doc, "First LNG is still on track for the second half of 2027. We're going to keep that commitment.")

    add_heading(doc, "Safety - we're doing well, and we can do better", level=3)
    add_body(doc,
        "Our TRIR for the quarter was 0.42, better than the industry "
        "benchmark of 0.65. That's great, and we want to keep it that way.")
    add_body(doc,
        "We had a hot-work near-miss on April 14 - HW-2026-0143. Nobody "
        "was hurt, but a fire watch was relieved without a formal handoff. "
        "We've already closed that gap in our procedure. The reminder: "
        "every permit, every handoff, every time. If something doesn't "
        "feel right, stop work.")

    add_heading(doc, "What's next - Phase 2", level=3)
    add_body(doc,
        "Our Commercial team is advancing three long-term SPAs with "
        "investment-grade counterparties. Our FERC authorization is in "
        "place. We're targeting a Phase 2 Final Investment Decision in "
        "the second half of this year. That's a big deal for this company "
        "and for the Rio Grande Valley.")

    add_heading(doc, "People news", level=3)
    add_bullet(doc, "We welcomed 47 new teammates in Q1 - introduce yourself.")
    add_bullet(doc, "Two teammates hit their 5-year anniversary this month.")
    add_bullet(doc, "Benefits open enrollment kicks off in May.")
    add_bullet(doc, "The next NEXT Digest drops Friday.")

    add_body(doc,
        "If you've got questions, send them to allhands@next-decade.com "
        "before the town hall on the 18th. See you then.")

    doc.add_paragraph()
    sig = doc.add_paragraph()
    add_runtext(sig, "Matt", size=11, color=BLACK)
    sig2 = doc.add_paragraph()
    add_runtext(sig2, "Matt Schatzman", bold=True, size=10, color=NAVY)
    add_runtext(sig2, "\nChairman and CEO, NextDecade Corporation", size=10, color=GREY)

    set_footer(doc, classification=True, mission=True, page_numbers=False)
    doc.save(str(output))
    print(f"Wrote {output}")


def build_next_digest_newsletter(output: Path):
    doc = new_document()

    # Masthead
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(m, "NEXT DIGEST", bold=True, size=36, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(sub, "The monthly newsletter for NextDecade teammates", italic=True, size=11, color=ORANGE)

    vol = doc.add_paragraph()
    vol.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(vol, "Volume 3  |  Issue 4  |  April 2026", size=10, color=GREY)

    add_horizontal_rule(doc, color=ORANGE)

    # Lead story
    add_heading(doc, "Lead Story: 68% and climbing", level=2)
    add_body(doc,
        "This quarter we crossed 68% complete on Phase 1 construction - a "
        "big step toward first LNG in the second half of 2027. We caught up "
        "with Brian Smith, COO, for a two-minute debrief.")
    add_body(doc,
        "\"The team executed,\" Brian said. \"Cold box set on Train 1, "
        "cryogenic piping installation starting, and procurement basically "
        "done. That's what 68% looks like.\"")

    # Safety corner
    add_heading(doc, "Safety Corner: April's hot-work reminder", level=2)
    add_body(doc,
        "On April 14 we had a near-miss on a hot-work permit - HW-2026-0143. "
        "A fire watch was relieved without a formal handoff. No injuries, "
        "but it's a lesson worth repeating: every permit, every handoff, "
        "every time. Procedure has been updated and a site-wide stand-down "
        "was held. Full details in Workplace under #hsse.")

    # People story — day in the life
    add_heading(doc, "Day in the Life: Maria Gonzalez, QA/QC Inspector", level=2)
    add_body(doc,
        "Maria joined NextDecade in 2024 after 12 years in oil and gas "
        "QA/QC in South Texas. She starts her day at 5:45 a.m. with a "
        "pre-shift stretch and a walkthrough of the cold-box area. "
        "\"Inspection is never the exciting part of the story,\" she "
        "laughed, \"but it's the reason Train 1 went up without a hitch. "
        "That's the part I'm proud of.\" When she's not on site, Maria "
        "coaches her daughter's soccer team in Los Fresnos.")

    # Rio Grande Valley
    add_heading(doc, "In the RGV: bilingual open house at Los Fresnos", level=2)
    add_body(doc,
        "On April 22, NextDecade is hosting an open house at the Los "
        "Fresnos Community Center from 6:00 p.m. to 8:00 p.m. Topics in "
        "English and Spanish: construction progress, hiring, and our "
        "community fund. RSVP at community@next-decade.com.")

    # CCS update
    add_heading(doc, "NEXT Carbon Solutions: Class VI permits", level=2)
    add_body(doc,
        "Our NEXT Carbon Solutions team advanced the Class VI permit "
        "program this quarter. The CCS project is designed to "
        "permanently store carbon dioxide associated with our LNG "
        "liquefaction - central to our lower carbon intensive energy mission.")

    # Events & benefits
    add_heading(doc, "What's coming up", level=2)
    add_bullet(doc, "April 18 - Q1 All-Hands town hall (Teams + auditorium)")
    add_bullet(doc, "April 22 - Bilingual open house, Los Fresnos")
    add_bullet(doc, "May 5-16 - Benefits open enrollment")
    add_bullet(doc, "May 8 - Next Operating Committee meeting")
    add_bullet(doc, "May 15 - Submit your \"Day in the Life\" idea to comms@next-decade.com")

    # Team shout-outs
    add_heading(doc, "Shout-outs", level=2)
    add_bullet(doc, "Commissioning team - cold-box set on Train 1, zero incidents.")
    add_bullet(doc, "Supply Chain - you kept the critical path clear through winter.")
    add_bullet(doc, "HSSE Advisors - you caught the hot-work gap early and we're safer for it.")

    # Tagline close
    add_horizontal_rule(doc, color=ORANGE)
    close = doc.add_paragraph()
    close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(close, "Delivering Energy for What's NEXT", bold=True, italic=True,
                size=14, color=NAVY)

    set_footer(doc, classification=True, mission=True, page_numbers=True)
    doc.save(str(output))
    print(f"Wrote {output}")


if __name__ == "__main__":
    out_dir = Path(sys.argv[1])
    out_dir.mkdir(parents=True, exist_ok=True)
    build_all_hands_email(out_dir / "All-Hands Email - Q1 2026.docx")
    build_next_digest_newsletter(out_dir / "NEXT Digest - April 2026.docx")
