"""Press Release — formal external register per Writing Style Guide §5."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from nextdecade_brand import (
    new_document, add_heading, add_body, add_bullet, add_runtext,
    set_footer, add_horizontal_rule,
    NAVY, BLACK, ORANGE, GREY,
    LEGAL_NAME, TICKER, HQ_ADDRESS, WEBSITE,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


FLS_LANGUAGE = (
    "This press release contains forward-looking statements within the meaning "
    "of Section 27A of the Securities Act of 1933 and Section 21E of the "
    "Securities Exchange Act of 1934. Forward-looking statements include "
    "statements with respect to NextDecade's plans, expectations, beliefs, "
    "intentions, and strategies regarding the Rio Grande LNG facility, "
    "NEXT Carbon Solutions, future contracting activity, final investment "
    "decisions, regulatory matters, construction and commissioning schedules, "
    "and financial results. The words 'anticipate,' 'assume,' 'budget,' "
    "'contemplate,' 'believe,' 'could,' 'estimate,' 'expect,' 'forecast,' "
    "'intend,' 'may,' 'plan,' 'potential,' 'project,' 'should,' 'target,' "
    "'will,' and similar expressions identify forward-looking statements. "
    "Actual results could differ materially from those projected. "
    "Factors that could cause actual results to differ include, among others, "
    "the ability to obtain financing for future phases, construction and "
    "operational risks at the Rio Grande LNG facility and associated pipeline, "
    "evolution of the global LNG market, the progress of NEXT Carbon Solutions "
    "carbon capture and storage projects, and changes in regulatory posture. "
    "NextDecade undertakes no obligation to publicly correct or update any "
    "forward-looking statement, whether as a result of new information, "
    "future events, or otherwise."
)


def build(output: str):
    doc = new_document()

    # Release tag line (tight, centered)
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(p0, "NEWS RELEASE", bold=True, size=11, color=ORANGE)

    # For-immediate-release + contact
    p1 = doc.add_paragraph()
    add_runtext(p1, "FOR IMMEDIATE RELEASE", bold=True, size=10, color=NAVY)
    add_runtext(p1, "\t\t\t\t", size=10, color=BLACK)
    add_runtext(p1, "April 15, 2026", size=10, color=BLACK)

    add_horizontal_rule(doc, color=NAVY)

    # Headline
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runtext(h,
        "NextDecade Reaches 68% Completion on Rio Grande LNG Phase 1 and "
        "Confirms Second-Half 2027 First LNG Target",
        bold=True, size=18, color=NAVY)

    # Subhead
    sh = doc.add_paragraph()
    add_runtext(sh,
        "Company targets a Final Investment Decision on Phase 2 in the "
        "second half of 2026; three anchor offtakers in advanced SPA negotiations.",
        italic=True, size=12, color=GREY)

    # Dateline + lede
    lede = doc.add_paragraph()
    add_runtext(lede, "HOUSTON, April 15, 2026 - ", bold=True, size=11, color=BLACK)
    add_runtext(lede,
        f"{LEGAL_NAME} ({TICKER}) today reported that construction of the "
        "Phase 1 Rio Grande LNG facility at the Port of Brownsville, Texas "
        "reached 68 percent completion through March 31, 2026. The Company "
        "continues to target first LNG from Train 1 in the second half of 2027. "
        "Trains 2 and 3 remain on the schedule presented at the Company's 2025 "
        "investor day.",
        size=11, color=BLACK)

    # Quote block
    q1 = doc.add_paragraph()
    add_runtext(q1, "\u201C", size=11, color=NAVY)
    add_runtext(q1,
        "Our teams are executing on one of the most complex LNG construction "
        "programs in the world, and they are doing it safely and on schedule. "
        "The next milestone is confirming the commercial and financing "
        "foundation for Phase 2, which we expect to bring to our Board for "
        "final investment decision before year-end.",
        italic=True, size=11, color=BLACK)
    add_runtext(q1, "\u201D", size=11, color=NAVY)
    attr = doc.add_paragraph()
    add_runtext(attr, "- Matt Schatzman, Chairman and Chief Executive Officer",
                bold=True, size=10, color=NAVY)

    # Body paragraphs
    add_body(doc,
        "Engineering is 99 percent complete and procurement is 88 percent "
        "complete, with all long-lead equipment on site or in transit. "
        "Commissioning of the non-cryogenic utilities is scheduled to begin "
        "in the third quarter of 2026.")

    add_body(doc,
        "Long-term sale and purchase agreements for Phase 2 volumes are "
        "progressing with three investment-grade counterparties. The Federal "
        "Energy Regulatory Commission authorization for Phase 2 remains in "
        "effect, and the Department of Energy non-Free Trade Agreement "
        "export authorization is expected to be confirmed in the second "
        "quarter of 2026.")

    add_body(doc,
        "The Rio Grande LNG site recorded 3.1 million construction hours in "
        "the first quarter of 2026, with a Total Recordable Incident Rate of "
        "0.42, below the industry benchmark. NextDecade's commitment to the "
        "workforce and to the Rio Grande Valley community remains central to "
        "the Company's execution.")

    # About
    add_heading(doc, f"About {LEGAL_NAME}", level=3)
    add_body(doc,
        f"{LEGAL_NAME} is an LNG development company. The Company is "
        "developing and constructing the Rio Grande LNG facility at the "
        "Port of Brownsville, Texas, one of the largest LNG export "
        "facilities in the United States. Through its subsidiary, "
        "NEXT Carbon Solutions, NextDecade is developing carbon capture "
        "and storage projects that seek to permanently store carbon "
        "dioxide associated with natural gas liquefaction. NextDecade is "
        "committed to providing the world access to lower carbon intensive "
        "energy. For more information, please visit "
        f"{WEBSITE}.")

    # Forward-looking statements
    add_heading(doc, "Forward-Looking Statements", level=3)
    flss = doc.add_paragraph()
    add_runtext(flss, FLS_LANGUAGE, size=8, color=GREY)

    # Contacts
    add_heading(doc, "Investor and Media Contacts", level=3)
    add_body(doc, "Investors:")
    add_runtext(doc.paragraphs[-1], "  ir@next-decade.com", bold=True, size=11, color=NAVY)
    add_body(doc, "Media:")
    add_runtext(doc.paragraphs[-1], "  media@next-decade.com", bold=True, size=11, color=NAVY)

    # HQ address block
    hq = doc.add_paragraph()
    hq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(hq, HQ_ADDRESS, size=8, color=GREY)

    # "###" release-end marker
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(end, "###", bold=True, size=11, color=NAVY)

    set_footer(doc, classification=False, mission=True, page_numbers=True)

    doc.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    build(sys.argv[1])
