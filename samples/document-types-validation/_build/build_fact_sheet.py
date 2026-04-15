"""Rio Grande LNG Fact Sheet — 1-page public-facing DOCX."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from nextdecade_brand import (
    new_document, add_heading, add_body, add_runtext,
    set_footer, add_horizontal_rule,
    NAVY, BLACK, ORANGE, GREY, LEGAL_NAME, WEBSITE, TICKER,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build(output: str):
    doc = new_document()

    # Hero title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(title, "RIO GRANDE LNG", bold=True, size=28, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(sub, "Fact Sheet  -  April 2026", italic=True, size=12, color=ORANGE)

    add_horizontal_rule(doc, color=ORANGE)

    # Two-column-like layout using a 2x1 table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(3.25)
        row.cells[1].width = Inches(3.25)
    left, right = table.rows[0].cells

    # Left column — project summary
    def lp(text, bold=False, size=10, color=BLACK):
        p = left.add_paragraph()
        add_runtext(p, text, bold=bold, size=size, color=color)

    add_runtext(left.paragraphs[0], "Project Overview", bold=True, size=14, color=NAVY)
    lp("Rio Grande LNG is an LNG liquefaction and export facility under "
       "construction at the Port of Brownsville, in the Rio Grande Valley "
       "of Texas. Phase 1 is comprised of three liquefaction trains with a "
       "combined nameplate capacity of 17.6 MTPA. Commercial operations for "
       "Train 1 are expected in the second half of 2027.")

    lp("Location", bold=True, size=12, color=NAVY)
    lp("Port of Brownsville, Cameron County, Texas")

    lp("Phase 1 Capacity", bold=True, size=12, color=NAVY)
    lp("3 Trains  |  17.6 MTPA nameplate")

    lp("First LNG", bold=True, size=12, color=NAVY)
    lp("Second half of 2027 (Train 1)")

    lp("Phase 2", bold=True, size=12, color=NAVY)
    lp("Up to 2 additional trains; FID targeted 2H 2026.")

    lp("Regulatory", bold=True, size=12, color=NAVY)
    lp("FERC authorized; DOE FTA and non-FTA authorizations in place.")

    # Right column — key facts
    def rp(text, bold=False, size=10, color=BLACK):
        p = right.add_paragraph()
        add_runtext(p, text, bold=bold, size=size, color=color)

    add_runtext(right.paragraphs[0], "At a Glance", bold=True, size=14, color=NAVY)

    rp("Parent", bold=True, size=12, color=NAVY)
    rp(f"{LEGAL_NAME} ({TICKER})")

    rp("Stock Listing", bold=True, size=12, color=NAVY)
    rp("NASDAQ: NEXT")

    rp("Project Milestones", bold=True, size=12, color=NAVY)
    rp(" - FID Phase 1: 2023")
    rp(" - Notice to Proceed: 2023")
    rp(" - Construction 68% complete (Q1 2026)")
    rp(" - First LNG: 2H 2027")

    rp("Community Commitment", bold=True, size=12, color=NAVY)
    rp("NextDecade is committed to the Rio Grande Valley. The Company "
       "supports local workforce development, bilingual community "
       "engagement, and environmental stewardship consistent with its "
       "FERC commitments.")

    rp("Carbon Reduction Pathway", bold=True, size=12, color=NAVY)
    rp("Through NEXT Carbon Solutions, NextDecade is developing carbon "
       "capture and storage projects designed to permanently store carbon "
       "dioxide from natural gas liquefaction.")

    # Bottom section — contacts and website
    doc.add_paragraph()
    add_horizontal_rule(doc, color=NAVY)

    contacts = doc.add_paragraph()
    contacts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runtext(contacts, f"{WEBSITE}   |   ir@next-decade.com   |   media@next-decade.com",
                bold=True, size=10, color=NAVY)

    set_footer(doc, classification=False, mission=True, page_numbers=False)

    doc.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    build(sys.argv[1])
