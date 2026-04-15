"""NextDecade brand-chrome helpers for python-docx.

All values sourced from extracted-specs.md and gap-report.md.

Per brand book / writing style guide:
- Body font: Segoe UI 11pt, black
- Headlines: Segoe UI 20pt, Blue Accent 1 Darker 50% (rendered as brand navy #002060)
- Subheads:  Segoe UI 16pt, brand navy
- Primary navy   #002060
- Primary orange #FC7134
- Primary green  #00B050

- Classification stamp (working default, per gap-report.md):
  "Confidential and Proprietary - This document is intended solely for internal
  use. Unauthorized disclosure, distribution, or reproduction is strictly
  prohibited. Content is preliminary and subject to revision."

- Footer mission line (canonical, spelled out):
  "NextDecade Corporation (NextDecade) is committed to providing the world access
  to lower carbon intensive energy through Rio Grande LNG & NEXT Carbon Solutions."
"""
from __future__ import annotations
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand constants
NAVY = RGBColor(0x00, 0x20, 0x60)
ORANGE = RGBColor(0xFC, 0x71, 0x34)
GREEN = RGBColor(0x00, 0xB0, 0x50)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x59, 0x59, 0x59)

FONT = "Segoe UI"

TAGLINE = "Delivering Energy for What's NEXT"
LEGAL_NAME = "NextDecade Corporation"
SHORT_NAME = "NextDecade"
TICKER = "NASDAQ: NEXT"
HQ_ADDRESS = "1000 Louisiana Street, Suite 3900, Houston, Texas 77002 USA"
WEBSITE = "www.next-decade.com"
MISSION_FOOTER = (
    f"{LEGAL_NAME} ({SHORT_NAME}) is committed to providing the world access "
    f"to lower carbon intensive energy through Rio Grande LNG "
    f"& NEXT Carbon Solutions."
)
CLASSIFICATION_FOOTER = (
    "Confidential and Proprietary - This document is intended solely for "
    "internal use. Unauthorized disclosure, distribution, or reproduction "
    "is strictly prohibited. Content is preliminary and subject to revision."
)


def new_document() -> Document:
    """Returns a new Document with Segoe UI 11pt body baseline."""
    doc = Document()
    # Page: Letter, 1in margins all sides, 0.5in header/footer
    for s in doc.sections:
        s.page_height = Inches(11)
        s.page_width = Inches(8.5)
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        s.header_distance = Inches(0.5)
        s.footer_distance = Inches(0.5)

    # Normal style -> Segoe UI 11pt black
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    _force_east_asian_font(normal, FONT)

    # Headings
    _style(doc, "Heading 1", size=20, bold=True, color=NAVY)
    _style(doc, "Heading 2", size=16, bold=True, color=NAVY)
    _style(doc, "Heading 3", size=13, bold=True, color=NAVY)
    return doc


def _style(doc, name: str, size: int, bold: bool, color: RGBColor):
    try:
        s = doc.styles[name]
    except KeyError:
        return
    s.font.name = FONT
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = color
    _force_east_asian_font(s, FONT)


def _force_east_asian_font(style, name: str):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def add_runtext(paragraph, text: str, *, bold=False, italic=False, size=None,
                color=None, font=FONT):
    run = paragraph.add_run(text)
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    # Force eastAsia font on the run to prevent Word from falling back
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:cs"), font)
    return run


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    add_runtext(
        p, text,
        bold=True,
        size={1: 20, 2: 16, 3: 13}.get(level, 13),
        color=NAVY,
    )
    return p


def add_body(doc, text: str, *, bold=False, italic=False):
    p = doc.add_paragraph()
    add_runtext(p, text, bold=bold, italic=italic, size=11, color=BLACK)
    return p


def add_bullet(doc, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        add_runtext(p, bold_lead, bold=True, size=11, color=BLACK)
    add_runtext(p, text, size=11, color=BLACK)
    return p


def set_footer(doc, classification: bool = True, mission: bool = True,
               page_numbers: bool = True):
    """Populate the first-section footer with NDLNG chrome."""
    sec = doc.sections[0]
    footer = sec.footer
    # Clear any default paragraph
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        if p.text:
            p.clear()
    p_main = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if mission:
        add_runtext(p_main, MISSION_FOOTER, italic=True, size=8, color=GREY)
    if classification:
        p_cls = footer.add_paragraph()
        p_cls.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runtext(p_cls, CLASSIFICATION_FOOTER, size=7, color=GREY)
    if page_numbers:
        p_pg = footer.add_paragraph()
        p_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runtext(p_pg, "Page ", size=8, color=GREY)
        _insert_page_field(p_pg)
        add_runtext(p_pg, " of ", size=8, color=GREY)
        _insert_pages_field(p_pg)


def set_header(doc, left_text: str, right_text: str = ""):
    sec = doc.sections[0]
    header = sec.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    # Left text
    add_runtext(p, left_text, bold=True, size=9, color=NAVY)
    if right_text:
        add_runtext(p, "\t" + right_text, size=9, color=GREY)
        # Add a right tab stop at the right margin (~6.5in from left margin)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Inches(6.5), alignment=2)  # 2 = RIGHT


def _insert_page_field(paragraph):
    _insert_field(paragraph, "PAGE")


def _insert_pages_field(paragraph):
    _insert_field(paragraph, "NUMPAGES")


def _insert_field(paragraph, field_code: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_horizontal_rule(doc, color: RGBColor = NAVY):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        add_runtext(p, h, bold=True, size=10, color=NAVY)
        _cell_shading(hdr[i], "E6EBF4")
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = ""
            p = row_cells[i].paragraphs[0]
            add_runtext(p, str(val), size=10, color=BLACK)
    return table


def _cell_shading(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
