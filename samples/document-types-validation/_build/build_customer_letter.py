"""Customer Letter — formal external voice with letterhead-style layout."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from nextdecade_brand import (
    new_document, add_body, add_runtext, set_footer,
    add_horizontal_rule,
    NAVY, BLACK, ORANGE, GREY,
    LEGAL_NAME, HQ_ADDRESS, WEBSITE,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build(output: str):
    doc = new_document()

    # Letterhead band (pseudo - since no real letterhead logo file is available,
    # we render a brand-chrome text block in the position the letterhead occupies)
    lh = doc.add_paragraph()
    add_runtext(lh, LEGAL_NAME, bold=True, size=16, color=NAVY)
    lh_sub = doc.add_paragraph()
    add_runtext(lh_sub, HQ_ADDRESS + "  |  " + WEBSITE, size=9, color=GREY)
    add_horizontal_rule(doc, color=ORANGE)

    # Date
    p_date = doc.add_paragraph()
    add_runtext(p_date, "April 15, 2026", size=11, color=BLACK)

    doc.add_paragraph()

    # Addressee block
    for line in [
        "Ms. Elena Rodriguez",
        "Chief Commercial Officer",
        "Valoria Energy Trading, Ltd.",
        "25 Canada Square",
        "London E14 5LQ",
        "United Kingdom",
    ]:
        p = doc.add_paragraph()
        add_runtext(p, line, size=11, color=BLACK)

    doc.add_paragraph()

    # Re: line
    p_re = doc.add_paragraph()
    add_runtext(p_re, "Re:  ", bold=True, size=11, color=NAVY)
    add_runtext(p_re,
        "Rio Grande LNG Phase 1 Construction Update and Phase 2 Offtake Discussion",
        bold=True, size=11, color=BLACK)

    doc.add_paragraph()

    # Salutation
    sal = doc.add_paragraph()
    add_runtext(sal, "Dear Ms. Rodriguez:", size=11, color=BLACK)

    # Body
    add_body(doc,
        "Thank you for the productive discussion last month between our "
        "respective commercial teams. Following that conversation, I am "
        "writing to provide a formal update on Rio Grande LNG Phase 1 "
        "construction and to confirm the path forward for Valoria's "
        "potential participation as a long-term Phase 2 offtaker.")

    add_body(doc,
        "As reported in our most recent public disclosures, Rio Grande LNG "
        "Phase 1 construction reached 68 percent completion as of March 31, "
        "2026. Engineering is 99 percent complete, procurement is 88 percent "
        "complete, and all long-lead equipment is on site or in transit. "
        "First LNG from Train 1 remains targeted for the second half of 2027, "
        "consistent with the schedule we have communicated at our investor events.")

    add_body(doc,
        "With respect to Phase 2, NextDecade is advancing the commercial "
        "and financing foundation necessary for a Final Investment Decision "
        "in the second half of 2026. Our offtake strategy for Phase 2 "
        "prioritizes long-term, investment-grade counterparties seeking "
        "diversified U.S. Gulf Coast supply. Based on our conversations, "
        "we believe Valoria is well-positioned to be one of those counterparties.")

    add_body(doc,
        "We propose to advance to a term sheet based on the structure our "
        "teams discussed: a 20-year SPA with FOB delivery at the Port of "
        "Brownsville, annual contract quantity to be confirmed, and a "
        "Henry Hub-linked pricing mechanism with agreed-upon fixed "
        "components. Our commercial team will circulate a working draft "
        "within two weeks of today. We would welcome the opportunity to "
        "host your team for a site visit during the second quarter to view "
        "the facility in its current state.")

    add_body(doc,
        "NextDecade appreciates the confidence Valoria has placed in our "
        "Rio Grande LNG project, and we look forward to taking the next "
        "step together. Please contact me directly if there is anything "
        "further I can provide.")

    # Closing
    doc.add_paragraph()
    close = doc.add_paragraph()
    add_runtext(close, "Sincerely,", size=11, color=BLACK)
    doc.add_paragraph()
    doc.add_paragraph()
    nm = doc.add_paragraph()
    add_runtext(nm, "Trent Williams", bold=True, size=11, color=NAVY)
    tt = doc.add_paragraph()
    add_runtext(tt, "Vice President, Commercial", size=11, color=BLACK)
    co = doc.add_paragraph()
    add_runtext(co, LEGAL_NAME, size=11, color=BLACK)

    # Cc
    doc.add_paragraph()
    cc = doc.add_paragraph()
    add_runtext(cc, "cc:  ", bold=True, size=10, color=NAVY)
    add_runtext(cc, "Matt Schatzman, Chairman and CEO, NextDecade Corporation",
                size=10, color=BLACK)

    set_footer(doc, classification=False, mission=True, page_numbers=True)
    doc.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    build(sys.argv[1])
