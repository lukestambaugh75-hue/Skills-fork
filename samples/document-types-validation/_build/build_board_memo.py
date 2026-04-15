"""Board Memo — formal external register per Writing Style Guide §5."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from nextdecade_brand import (
    new_document, add_heading, add_body, add_bullet, add_runtext,
    set_footer, set_header, add_horizontal_rule, NAVY, BLACK, ORANGE,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build(output: str):
    doc = new_document()

    # Header chrome
    set_header(doc, left_text="NextDecade Corporation", right_text="Board Memorandum")

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runtext(title, "BOARD MEMORANDUM", bold=True, size=20, color=NAVY)

    add_horizontal_rule(doc, color=ORANGE)

    # To / From / Date / Subject block (classic memo header)
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.1)
    table.columns[1].width = Inches(5.4)
    rows = [
        ("TO:", "Board of Directors, NextDecade Corporation"),
        ("FROM:", "Matt Schatzman, Chairman and Chief Executive Officer"),
        ("DATE:", "April 15, 2026"),
        ("SUBJECT:", "Q1 2026 Rio Grande LNG Construction Progress and Phase 2 FID Readiness"),
    ]
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        add_runtext(c0.paragraphs[0], label, bold=True, size=11, color=NAVY)
        c1.text = ""
        add_runtext(c1.paragraphs[0], value, size=11, color=BLACK)

    add_horizontal_rule(doc, color=NAVY)

    # Executive summary
    add_heading(doc, "Executive Summary", level=2)
    add_body(doc,
        "Rio Grande LNG Phase 1 construction advanced to 68% complete through "
        "March 31, 2026, consistent with the schedule presented to the Board "
        "at the February 2026 meeting. First LNG remains targeted for the "
        "second half of 2027. The Phase 2 Final Investment Decision package "
        "is on track for Board consideration in the third quarter of 2026. "
        "No material safety events occurred during the quarter."
    )

    # Detail sections
    add_heading(doc, "Construction Status", level=2)
    add_body(doc,
        "Engineering is 99% complete. Procurement is 88% complete with all "
        "long-lead equipment on site or in transit. Construction is 52% "
        "complete across Trains 1, 2, and 3. The liquefaction cold box for "
        "Train 1 was set during the quarter without incident; commissioning "
        "of the non-cryogenic utilities is scheduled to begin in July 2026."
    )

    add_heading(doc, "Commercial and Regulatory", level=2)
    add_body(doc,
        "Long-term SPAs for Phase 2 volumes are progressing with three "
        "investment-grade counterparties. The FERC authorization for "
        "Phase 2 remains in effect; the Department of Energy non-FTA export "
        "authorization is expected to be confirmed in the second quarter of "
        "2026. There have been no changes in the regulatory posture "
        "previously disclosed to the Board."
    )

    add_heading(doc, "Financial", level=2)
    add_body(doc,
        "Phase 1 remains on budget at $18.4 billion. As of March 31, 2026, "
        "total project spend was $9.8 billion, funded from the project "
        "debt facility ($6.2B), committed equity ($3.0B), and the corporate "
        "revolver ($0.6B). The Company's liquidity position provides "
        "adequate headroom through first LNG."
    )

    add_heading(doc, "Safety and HSSE", level=2)
    add_body(doc,
        "The site recorded 3.1 million construction hours in the quarter "
        "with a Total Recordable Incident Rate of 0.42, below the industry "
        "benchmark of 0.65. No Life-Saving Rule violations were recorded. "
        "One near-miss event related to hot work (HW-2026-0143, April 14) "
        "was investigated; corrective actions are complete and will be "
        "discussed in the HSSE Committee pre-read."
    )

    add_heading(doc, "Phase 2 Readiness", level=2)
    add_body(doc,
        "Management recommends the Board direct the preparation of a formal "
        "Phase 2 FID package for consideration at the August 2026 meeting. "
        "Items that must be resolved prior to FID include:"
    )
    add_bullet(doc, "Final SPA execution with the three anchor offtakers.")
    add_bullet(doc, "Phase 2 debt facility commitment letters.")
    add_bullet(doc, "Updated EPC cost certainty from the primary contractor.")
    add_bullet(doc, "Confirmation of Phase 2 incremental equity funding.")
    add_bullet(doc, "Board-level review of community and stakeholder engagement metrics in the Rio Grande Valley.")

    add_heading(doc, "Recommended Board Action", level=2)
    add_body(doc,
        "The Board is asked to (i) acknowledge receipt of this report, "
        "(ii) authorize management to prepare the Phase 2 FID package for "
        "the August 2026 meeting, and (iii) direct the HSSE Committee to "
        "review the HW-2026-0143 corrective action plan at its May 2026 session."
    )

    # Signature block
    doc.add_paragraph()
    sig = doc.add_paragraph()
    add_runtext(sig, "Respectfully submitted,", size=11, color=BLACK)
    doc.add_paragraph()
    doc.add_paragraph()
    name = doc.add_paragraph()
    add_runtext(name, "Matt Schatzman", bold=True, size=11, color=NAVY)
    title2 = doc.add_paragraph()
    add_runtext(title2, "Chairman and Chief Executive Officer, NextDecade Corporation",
                size=11, color=BLACK)

    set_footer(doc, classification=True, mission=True, page_numbers=True)

    doc.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    build(sys.argv[1])
