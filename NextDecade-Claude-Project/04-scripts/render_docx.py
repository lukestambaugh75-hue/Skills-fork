#!/usr/bin/env python3
"""docxtpl render pipeline for NextDecade governance documents.

Supports: procedure, standard, guidance. Each has a Jinja-tagged template and
a JSON schema under ../templates/. Pipeline:
  1. Lint the Jinja template (scripts/lint_docx_template.py).
  2. If clean/warnings only -> render with docxtpl + StrictUndefined.
  3. If errors -> fall back to walk-and-replace on the ORIGINAL template.
  4. Optional: convert the rendered .docx to .pdf via LibreOffice.

CLI:
  python render_docx.py <doc_type> <input.json> <output.docx> [--pdf] [--no-fallback]
    doc_type: procedure | standard | guidance

Module:
  from render_docx import render
  report = render("procedure", data_dict, "/path/to/output.docx", pdf=True)
"""
from __future__ import annotations
import sys, json, subprocess, shutil, os
from pathlib import Path
from typing import Any

HERE = Path(__file__).resolve().parent
# NOTE: TEMPLATES and UPLOADS path constants below differ between the two
# render_docx.py copies (skills/docx/scripts/ vs NextDecade-Claude-Project/
# 04-scripts/) because each copy resolves templates relative to its own
# directory. The smoke test strips path-constant lines before diffing so
# only functional drift is reported.
TEMPLATES = HERE.parent / "02-templates"
LINTER = HERE / "lint_docx_template.py"

# Template registry — extend here when adding new document types.
# `original_template` is used only by the walk-and-replace FALLBACK path; if it
# is missing (e.g., a template was renamed/swapped), the pipeline still works
# via docxtpl — the fallback just won't be available for that type.
# Resolve paths relative to this file: 04-scripts/ -> parent project dir
UPLOADS = HERE.parent / "03-original-templates"
DOC_TYPES = {
    "procedure": {
        # Single source of formatting: the Jinja template is used by BOTH the
        # happy-path (strict docxtpl) and the fallback (lenient docxtpl). The
        # 03-original-templates/ copy is no longer consulted for procedure —
        # the file on disk remains for archival only.
        "jinja_template": TEMPLATES / "Procedure Template (Jinja).docx",
        "schema": TEMPLATES / "procedure_schema.json",
        "fallback": "walk_replace_procedure",
    },
    "standard": {
        # Same single-source invariant as procedure above.
        "jinja_template": TEMPLATES / "Standard Template (Jinja).docx",
        "schema": TEMPLATES / "standard_schema.json",
        "fallback": "walk_replace_standard",
    },
    "guidance": {
        "jinja_template": TEMPLATES / "Guidance Template (Jinja).docx",
        "original_template_candidates": [UPLOADS / "Guidance Template.docx"],
        "schema": TEMPLATES / "guidance_schema.json",
        "fallback": "walk_replace_guidance",
    },
}

def _resolve_original(doc_type: str) -> Path | None:
    cfg = DOC_TYPES[doc_type]
    for cand in cfg.get("original_template_candidates", []):
        if cand.exists():
            return cand
    return None


def _strict_env():
    from jinja2 import Environment, StrictUndefined
    return Environment(undefined=StrictUndefined, autoescape=False)


def _xml_escape_data(obj):
    """Recursively escape XML-special characters (&, <, >) in string values.

    docxtpl injects rendered text directly into Word XML. If the data contains
    raw '&', '<', or '>' the XML parser fails with 'xmlParseEntityRef'. This
    pre-escaping makes the docxtpl path safe for arbitrary user content
    including ampersands, angle brackets, and em-dashes.
    """
    from xml.sax.saxutils import escape
    if isinstance(obj, str):
        return escape(obj)
    if isinstance(obj, dict):
        return {k: _xml_escape_data(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_xml_escape_data(v) for v in obj]
    return obj


def render_via_docxtpl(data: dict, template: Path, output: Path) -> None:
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(str(template))
    tpl.render(_xml_escape_data(data), jinja_env=_strict_env())
    tpl.save(str(output))


def render_via_docxtpl_lenient(data: dict, template: Path, output: Path) -> None:
    """Fallback renderer for procedure + standard.

    Reads from the SAME Jinja template as render_via_docxtpl so output
    formatting (theme, styles, fonts, margins, headers, footers) is guaranteed
    to match the happy path. The only difference is Undefined handling:
    missing keys / attrs render as empty strings instead of raising, which is
    what makes this a useful safety net when data is incomplete.

    Guidance keeps its own walk-and-replace implementation below; procedure
    and standard are both served by this function.
    """
    from docxtpl import DocxTemplate
    from jinja2 import Environment, ChainableUndefined
    env = Environment(undefined=ChainableUndefined, autoescape=False)
    tpl = DocxTemplate(str(template))
    tpl.render(_xml_escape_data(data), jinja_env=env)
    tpl.save(str(output))


def _post_render_cover_fixup(output: Path, data: dict, doc_type: str) -> list[str]:
    """Patch cover-page text boxes and headers in the rendered .docx.

    The Standard and Guidance Jinja templates have plain-text placeholders
    ("NAME", "xxx-xxx-xxx-xxx-xxx-#####") inside text boxes and headers that
    docxtpl cannot reach because they're not Jinja markers. This function
    does a targeted XML search-and-replace inside the .docx zip.

    Returns a list of patches applied (for logging).
    """
    import re, zipfile

    doc_name = data.get("document_name", "")
    doc_num = data.get("doc_number", "")
    if not doc_name and not doc_num:
        return []

    type_label = {"standard": "STANDARD", "guidance": "GUIDANCE"}.get(doc_type)
    if type_label is None:
        return []  # procedure has its own cover-page handling via docxtpl markers

    # Define replacements per doc type
    replacements = []
    if doc_name:
        # Cover text box: "NAME" (standalone or part of "NAME STANDARD DOCUMENT")
        replacements.append(("NAME", doc_name))
        # Header: "[Name] Standard" or "[GUIDANCE Document NAME]"
        replacements.append((f"[Name] {type_label.title()}", f"{doc_name}"))
        replacements.append((f"[{type_label} Document NAME]", f"{doc_name}"))
    if doc_num:
        # Cover text box uses "xxx-xxx-xxx-xxx-xxx-#####"; headers may use
        # "xxx-xxx-xxx-xxx-xxx-" (trailing dash, no #####). Replace both.
        replacements.append(("xxx-xxx-xxx-xxx-xxx-#####", doc_num))
        replacements.append(("Doc. No. xxx-xxx-xxx-xxx-xxx-", f"Doc. No. {doc_num}"))

    patches = []
    with zipfile.ZipFile(output) as zin:
        files = {n: zin.read(n) for n in zin.namelist()}
        infos = {n: zin.getinfo(n) for n in zin.namelist()}

    target_parts = [n for n in files if n.startswith("word/") and n.endswith(".xml")]
    for part_name in target_parts:
        content = files[part_name].decode("utf-8")
        changed = False
        for old, new in replacements:
            if old in content:
                content = content.replace(old, new)
                patches.append(f"{part_name}: '{old}' -> '{new}'")
                changed = True
        if changed:
            files[part_name] = content.encode("utf-8")

    if patches:
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data_bytes in files.items():
                zout.writestr(infos[name], data_bytes)

    return patches


def _remove_template_scaffolding(output: Path, doc_type: str) -> None:
    """Remove known template scaffolding from the rendered .docx.

    Procedure: removes "Use the following for notes, cautions, and warnings."
    paragraph and the Note/Caution/Warning demo table that follows.
    Standard/Guidance: removes placeholder boilerplate paragraphs
    (e.g., "[CONTENT TITLE]", "Enter text here") if any survive rendering.
    """
    from docx import Document

    doc = Document(str(output))
    changed = False

    if doc_type == "procedure":
        # Remove boilerplate instruction paragraph
        for p in list(doc.paragraphs):
            if "Use the following for notes, cautions, and warnings" in p.text:
                p._element.getparent().remove(p._element)
                changed = True

        # Remove Note/Caution/Warning demo table
        for t in list(doc.tables):
            cell_texts = []
            for row in t.rows:
                for cell in row.cells:
                    cell_texts.append(cell.text.strip())
            all_text = " ".join(cell_texts)
            if ("Note:" in all_text and "Caution:" in all_text
                    and "Warning:" in all_text
                    and "alerts users" in all_text):
                t._element.getparent().remove(t._element)
                changed = True

    elif doc_type in ("standard", "guidance"):
        scaffolding_markers = [
            "[CONTENT TITLE]",
            "Enter text here",
            "Click here to enter text",
        ]
        for p in list(doc.paragraphs):
            for marker in scaffolding_markers:
                if marker in p.text:
                    p._element.getparent().remove(p._element)
                    changed = True
                    break

    if changed:
        doc.save(str(output))


# ---------------------------------------------------------------------------
# Walk-and-replace fallback (per-doc-type)
# ---------------------------------------------------------------------------
def _wr_common_setup(original: Path, output: Path):
    from docx import Document
    from copy import deepcopy
    from docx.text.paragraph import Paragraph
    shutil.copy2(original, output)
    doc = Document(str(output))

    def _set(p, text):
        runs = p.runs
        if not runs: p.add_run(text); return
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)

    def _h1(token):
        for i, p in enumerate(doc.paragraphs):
            if p.style and p.style.name == "Heading 1" and token.lower() in p.text.strip().lower():
                return i
        return None

    def _fill_after(idx, text):
        if idx is None: return
        for k in range(idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[k]
            if p.style and p.style.name == "Heading 1": return
            if p.text.strip(): _set(p, text); return

    return doc, _set, _h1, _fill_after, deepcopy, Paragraph


def _wr_fill_content_sections(doc, data, _set, deepcopy, Paragraph):
    sections = data["content_sections"]
    slots = [i for i, p in enumerate(doc.paragraphs)
             if p.style and p.style.name == "Heading 1" and "[CONTENT TITLE]" in p.text]

    def fill(heading_idx, section):
        h = doc.paragraphs[heading_idx]
        _set(h, section["title"])
        intro_set = False
        bullet_paragraphs = []
        for k in range(heading_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[k]
            if p.style and p.style.name == "Heading 1": break
            if not p.text.strip(): continue
            if not intro_set:
                _set(p, section["intro"]); intro_set = True
            else:
                bullet_paragraphs.append(p)
        bullets = section["bullets"]
        if len(bullets) <= len(bullet_paragraphs):
            for i, b in enumerate(bullets):
                _set(bullet_paragraphs[i], b)
            for extra in bullet_paragraphs[len(bullets):]:
                extra._element.getparent().remove(extra._element)
        else:
            for i, p in enumerate(bullet_paragraphs):
                _set(p, bullets[i])
            last = bullet_paragraphs[-1]
            for b in bullets[len(bullet_paragraphs):]:
                new_el = deepcopy(last._element)
                last._element.addnext(new_el)
                last = Paragraph(new_el, last._parent)
                _set(last, b)

    n_slots, n_sections = len(slots), len(sections)
    if n_sections <= n_slots:
        for i in range(n_sections):
            slot_idx = [j for j, p in enumerate(doc.paragraphs)
                        if p.style and p.style.name == "Heading 1"
                        and "[CONTENT TITLE]" in p.text][0]
            fill(slot_idx, sections[i])
        while True:
            remaining = [j for j, p in enumerate(doc.paragraphs)
                         if p.style and p.style.name == "Heading 1"
                         and "[CONTENT TITLE]" in p.text]
            if not remaining: break
            s = remaining[0]
            h = doc.paragraphs[s]
            to_remove = [h._element]
            for k in range(s + 1, len(doc.paragraphs)):
                p = doc.paragraphs[k]
                if p.style and p.style.name == "Heading 1": break
                to_remove.append(p._element)
            for el in to_remove:
                el.getparent().remove(el)
    else:
        for i in range(n_slots):
            slot_idx = [j for j, p in enumerate(doc.paragraphs)
                        if p.style and p.style.name == "Heading 1"
                        and "[CONTENT TITLE]" in p.text][0]
            fill(slot_idx, sections[i])
        last_title = sections[n_slots - 1]["title"]
        last_h_idx = max(j for j, p in enumerate(doc.paragraphs)
                         if p.style and p.style.name == "Heading 1"
                         and p.text.strip() == last_title)
        end_idx = len(doc.paragraphs)
        for k in range(last_h_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[k]
            if p.style and p.style.name == "Heading 1":
                end_idx = k; break
        block_elements = [doc.paragraphs[i]._element for i in range(last_h_idx, end_idx)]
        insert_after = block_elements[-1]
        for sec in sections[n_slots:]:
            new_elements = [deepcopy(el) for el in block_elements]
            for el in new_elements:
                insert_after.addnext(el); insert_after = el
            new_h_idx = max(j for j, p in enumerate(doc.paragraphs)
                            if p.style and p.style.name == "Heading 1"
                            and p.text.strip() == last_title)
            fill(new_h_idx, sec)


def _wr_fill_tables(doc, table_plans, _set):
    for t in doc.tables:
        headers = tuple(c.text.strip() for c in t.rows[0].cells)
        for plan in table_plans:
            if headers[:len(plan["match"])] == plan["match"]:
                plan["apply"](t)
                break


def walk_replace_procedure(data: dict, output: Path):
    """Fallback for procedure. Renders via lenient docxtpl against the same
    Jinja template used by the happy path. See render_via_docxtpl_lenient for
    the why; the short version is: one template file, one source of
    formatting, drift impossible by construction.
    """
    cfg = DOC_TYPES["procedure"]
    render_via_docxtpl_lenient(data, cfg["jinja_template"], output)


def walk_replace_standard(data: dict, output: Path):
    """Fallback for standard. Same design as walk_replace_procedure above."""
    cfg = DOC_TYPES["standard"]
    render_via_docxtpl_lenient(data, cfg["jinja_template"], output)


def walk_replace_guidance(data: dict, output: Path):
    original = _resolve_original("guidance")
    if original is None: raise RuntimeError("No guidance template available for walk-and-replace fallback.")
    doc, _set, _h1, _fill_after, deepcopy, Paragraph = _wr_common_setup(original, output)
    _fill_after(_h1("PURPOSE"),                  data["purpose_text"])
    # Remove leftover template boilerplate paragraphs in PURPOSE section
    _boilerplate_markers = [
        "Best Practices Recommendations",  # covers smart-quote variants
        "The Guideline applies to all employees, officers, directors",
    ]
    for p in list(doc.paragraphs):
        for marker in _boilerplate_markers:
            if marker in p.text:
                p._element.getparent().remove(p._element)
                break
    _fill_after(_h1("INTEGRATED GOVERNANCE"),    data["governance_text"])
    # GUIDELINE: intro + bullets (no title repetition)
    gi = _h1("GUIDELINE")
    if gi is not None:
        intro_set = False
        bullet_paragraphs = []
        for k in range(gi + 1, len(doc.paragraphs)):
            p = doc.paragraphs[k]
            if p.style and p.style.name == "Heading 1": break
            if not p.text.strip(): continue
            if not intro_set:
                _set(p, data["guideline_intro"]); intro_set = True
            else:
                bullet_paragraphs.append(p)
        bullets = data["guideline_bullets"]
        if len(bullets) <= len(bullet_paragraphs):
            for i, b in enumerate(bullets):
                _set(bullet_paragraphs[i], b)
            for extra in bullet_paragraphs[len(bullets):]:
                extra._element.getparent().remove(extra._element)
        else:
            for i, p in enumerate(bullet_paragraphs):
                _set(p, bullets[i])
            last = bullet_paragraphs[-1]
            for b in bullets[len(bullet_paragraphs):]:
                new_el = deepcopy(last._element)
                last._element.addnext(new_el)
                last = Paragraph(new_el, last._parent)
                _set(last, b)

    _wr_fill_tables(doc, _table_plans_guidance(data, _set), _set)
    doc.save(str(output))


def _table_plans_guidance(data, _set):
    return _table_plans_common(data, _set)


def _table_plans_common(data, _set):
    def fill_raci(t):
        if len(t.rows) > 1:
            r = t.rows[1]
            raci = data["raci"]
            _set(r.cells[0].paragraphs[0], raci["responsible"])
            _set(r.cells[1].paragraphs[0], raci["accountable"])
            _set(r.cells[2].paragraphs[0], raci["consulted"])
            _set(r.cells[3].paragraphs[0], raci["informed"])

    def fill_approval(t):
        if len(t.rows) > 1:
            r = t.rows[1]
            ap = data["approval"]
            _set(r.cells[0].paragraphs[0], ap["issuer"])
            _set(r.cells[1].paragraphs[0], ap["adopted_by"])
            _set(r.cells[2].paragraphs[0], ap["effective_date"])

    def fill_rev(t):
        data_rows = list(t.rows)[1:]
        for i, rev in enumerate(data["revision_history"]):
            # Schema uses "rev" across all doc types. Accept legacy "number"
            # for back-compat with any still-unmigrated input JSONs.
            rev_label = rev.get("rev", rev.get("number", ""))
            if i < len(data_rows):
                r = data_rows[i]
                _set(r.cells[0].paragraphs[0], rev_label)
                _set(r.cells[1].paragraphs[0], rev["description"])
            else:
                new_row = t.add_row()
                new_row.cells[0].text = rev_label
                new_row.cells[1].text = rev["description"]
        for r in data_rows[len(data["revision_history"]):]:
            r._element.getparent().remove(r._element)

    return [
        {"match": ("Responsible", "Accountable", "Consulted", "Informed"), "apply": fill_raci},
        {"match": ("Issuer / Title", "Adopted By", "Effective / Amended Date"), "apply": fill_approval},
        {"match": ("Revision", "Description of Changes and Notes"), "apply": fill_rev},
    ]


FALLBACK_FNS = {
    "walk_replace_procedure": walk_replace_procedure,
    "walk_replace_standard":  walk_replace_standard,
    "walk_replace_guidance":  walk_replace_guidance,
}


# ---------------------------------------------------------------------------
# PDF export via LibreOffice
# ---------------------------------------------------------------------------
def convert_to_pdf(docx_path: Path, pdf_path: Path | None = None) -> Path:
    """Convert a .docx to .pdf via headless LibreOffice (soffice).

    soffice must be installed with the Writer component. On Debian/Ubuntu:
      apt-get install libreoffice-writer

    Uses an isolated -env:UserInstallation profile so concurrent soffice
    invocations don't collide on shared profile state (required in sandboxes).
    """
    import tempfile
    docx_path = Path(docx_path)
    if pdf_path is None:
        pdf_path = docx_path.with_suffix(".pdf")
    out_dir = pdf_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as tmpdir:
        profile_arg = f"-env:UserInstallation=file://{tmpdir}/lo_profile"
        cmd = ["soffice", "--headless", profile_arg, "--convert-to", "pdf",
               "--outdir", str(out_dir), str(docx_path)]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed (exit {result.returncode}):\n"
                f"stderr: {result.stderr}\nstdout: {result.stdout}"
            )
    actual = out_dir / (docx_path.stem + ".pdf")
    if actual != pdf_path and actual.exists():
        shutil.move(str(actual), str(pdf_path))
    if not pdf_path.exists():
        raise RuntimeError(f"PDF was not produced at {pdf_path}")
    return pdf_path


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------
def render(doc_type: str, data: dict, output: str | Path,
           allow_fallback: bool = True, pdf: bool = False) -> dict:
    if doc_type not in DOC_TYPES:
        raise ValueError(f"Unknown doc_type {doc_type!r}. "
                         f"Expected one of: {sorted(DOC_TYPES)}")
    cfg = DOC_TYPES[doc_type]
    output = Path(output)
    report = {"doc_type": doc_type, "output": str(output), "path": None, "warnings": []}

    # Step 1: lint
    proc = subprocess.run(
        [sys.executable, str(LINTER), str(cfg["jinja_template"]), str(cfg["schema"])],
        capture_output=True, text=True,
    )
    lint_report = json.loads(proc.stdout) if proc.stdout else {}
    lint_exit = proc.returncode
    report["lint"] = {
        "exit": lint_exit,
        "issue_count": len(lint_report.get("issues", [])),
        "warning_count": len(lint_report.get("warnings", [])),
    }

    # Step 2: route
    if lint_exit in (0, 2):
        try:
            render_via_docxtpl(data, cfg["jinja_template"], output)
            report["path"] = "docxtpl"
            if lint_exit == 2:
                report["warnings"] = [w["message"] for w in lint_report.get("warnings", [])]
        except Exception as e:
            if allow_fallback:
                FALLBACK_FNS[cfg["fallback"]](data, output)
                report["path"] = "walk_and_replace_fallback"
                report["warnings"].append(f"docxtpl raised: {e}; used fallback")
            else:
                raise
    else:
        if allow_fallback:
            FALLBACK_FNS[cfg["fallback"]](data, output)
            report["path"] = "walk_and_replace_fallback"
            report["warnings"].append(
                "Jinja template has unrecoverable issues; rendered via walk-and-replace. "
                "Fix the Jinja template to restore the fast path."
            )
            report["lint_issues"] = lint_report.get("issues", [])
        else:
            raise RuntimeError(
                f"Template lint failed with exit {lint_exit}; fallback disabled. "
                f"Issues: {lint_report.get('issues')}"
            )

    # Step 2.5a: post-render cover-page fixup (Standard/Guidance text boxes + headers)
    try:
        cover_patches = _post_render_cover_fixup(output, data, doc_type)
        if cover_patches:
            report["cover_patches"] = cover_patches
    except Exception as e:
        report["warnings"].append(f"Cover-page fixup failed: {e}")

    # Step 2.5b: remove template scaffolding (applies to all doc types)
    try:
        _remove_template_scaffolding(output, doc_type)
    except Exception as e:
        report["warnings"].append(f"Scaffolding cleanup failed: {e}")

    # Step 3: optional PDF
    if pdf:
        try:
            pdf_path = convert_to_pdf(output)
            report["pdf"] = str(pdf_path)
        except Exception as e:
            report["warnings"].append(f"PDF export failed: {e}")
    return report


# Back-compat shim for the original entry point
def render_procedure(data: dict, output, allow_fallback: bool = True) -> dict:
    return render("procedure", data, output, allow_fallback=allow_fallback)


def main():
    if len(sys.argv) < 4:
        print("Usage: render_docx.py <doc_type> <input.json> <output.docx> "
              "[--pdf] [--no-fallback]", file=sys.stderr)
        print(f"  doc_type one of: {sorted(DOC_TYPES)}", file=sys.stderr)
        sys.exit(64)
    doc_type = sys.argv[1]
    data = json.loads(Path(sys.argv[2]).read_text())
    output = sys.argv[3]
    pdf = "--pdf" in sys.argv
    allow_fallback = "--no-fallback" not in sys.argv
    report = render(doc_type, data, output, allow_fallback=allow_fallback, pdf=pdf)
    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
