#!/usr/bin/env python3
"""Extract content and images from an existing .docx into schema-compatible JSON.

Reads a NextDecade governance document (Procedure, Standard, or Guidance)
and produces:
  - A JSON file matching the render_docx.render() input schema
  - An images/ directory with any embedded images extracted verbatim

Auto-detects doc_type from section headings. Works on both labeled and
unlabeled documents (no pre-existing JSON required).

CLI:
  python extract_docx.py <input.docx> [output.json] [--images-dir DIR]
    Outputs JSON to stdout if output.json is not given.
    --images-dir defaults to a sibling 'images/' dir next to output.json.

Module:
  from extract_docx import extract
  result = extract("/path/to/doc.docx", "/path/to/out.json",
                   images_dir="/path/to/imgs/")
  # result: {"doc_type": "procedure", "output_json": "...", "images": [...]}
"""
from __future__ import annotations
import sys, json, re, zipfile
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn as _qn
except ImportError:
    Document = None  # type: ignore


# ---------------------------------------------------------------------------
# Heading-to-field mapping tables
# ---------------------------------------------------------------------------

# Procedure: numbered headings like "1.0 PURPOSE", "2.0 SCOPE", etc.
_PROC_HEADING_MAP = {
    "PURPOSE": "purpose_text",
    "SCOPE": "scope_text",
    "ROLES AND RESPONSIBILITIES": "roles",
    "PERSONAL PROTECTIVE EQUIPMENT": "ppe_paragraphs",
    "RECORD KEEPING": "recordkeeping_text",
    "DEFINITIONS": "definitions",
    "ABBREVIATIONS": "abbreviations",
    "REFERENCES": "references",
    "APPENDICES": "appendices",
}

# Standard: section headings
_STD_HEADING_MAP = {
    "INTRODUCTION": "introduction_text",
    "SCOPE": "scope_text",
    "INTEGRATED GOVERNANCE FRAMEWORK": "governance_text",
    "DEFINITIONS": "definitions",
    "REFERENCES": "references",
    "EXCEPTION REQUEST": "exception_text",
    "CONTINUOUS IMPROVEMENT": "continuous_text",
    "OWNERSHIP": "_ownership",
    "APPROVAL": "_approval",
    "REVISION HISTORY": "_revision_history",
}

# Guidance: section headings
_GDN_HEADING_MAP = {
    "PURPOSE": "purpose_text",
    "INTEGRATED GOVERNANCE FRAMEWORK": "governance_text",
    "GUIDELINE": "_guideline",
    "OWNERSHIP": "_ownership",
    "APPROVAL": "_approval",
    "REVISION HISTORY": "_revision_history",
}


def _is_heading(p) -> bool:
    """Return True if paragraph uses any top-level heading style.
    Works on both python-docx Paragraph objects and raw lxml CT_P elements.
    """
    try:
        style_name = (p.style.name or "") if p.style else ""
    except AttributeError:
        # Fallback: read w:pStyle directly from XML when p is a detached element
        # (Paragraph(el, CT_Body) lacks .part so p.style raises AttributeError)
        try:
            pPr = p._element.find(_qn("w:pPr")) if hasattr(p, "_element") else p.find(_qn("w:pPr"))
            pStyle = pPr.find(_qn("w:pStyle")) if pPr is not None else None
            style_name = (pStyle.get(_qn("w:val"), "") if pStyle is not None else "") or ""
        except Exception:
            return False
    return (
        "Heading" in style_name
        or "RGLNG" in style_name      # NextDecade procedure heading styles
        or style_name.lower().startswith("hdg")
    )


def _normalize_heading(text: str) -> str:
    """Strip numbering prefix (e.g. '1.0 ') and normalize whitespace."""
    text = re.sub(r"^\d+(\.\d+)*\s+", "", text.strip())
    return text.upper()


def _detect_doc_type(paragraphs) -> str:
    """Infer document type from heading patterns."""
    headings = set()
    for p in paragraphs:
        if _is_heading(p):
            headings.add(_normalize_heading(p.text))

    if "GUIDELINE" in headings:
        return "guidance"
    if "INTRODUCTION" in headings:
        return "standard"
    # Procedure uses numbered headings — look for "1.0" prefix pattern
    for p in paragraphs:
        if _is_heading(p) and re.match(r"\d+\.\d+", p.text.strip()):
            return "procedure"
    # Fallback: if PURPOSE but no GUIDELINE/INTRODUCTION → procedure
    if "PURPOSE" in headings:
        return "procedure"
    return "procedure"  # default


# ---------------------------------------------------------------------------
# Paragraph collection helpers
# ---------------------------------------------------------------------------

def _collect_text_until_next_heading(paragraphs, start_idx: int) -> tuple[str, int]:
    """Gather body text after paragraphs[start_idx] until the next Heading paragraph.
    Returns (combined_text, next_heading_index).
    """
    texts = []
    i = start_idx + 1
    while i < len(paragraphs):
        p = paragraphs[i]
        if _is_heading(p):
            break
        texts.append(p.text)
        i += 1
    return "\n\n".join(t for t in texts if t.strip()), i


def _collect_bullets_until_next_heading(paragraphs, start_idx: int) -> tuple[str, list[str], int]:
    """Collect intro paragraph + bullets from a list section.
    First non-empty body paragraph is the intro; subsequent List-style
    paragraphs are bullets. Returns (intro, bullets, next_heading_index).
    """
    intro = ""
    bullets = []
    i = start_idx + 1
    while i < len(paragraphs):
        p = paragraphs[i]
        if _is_heading(p):
            break
        style_name = (p.style.name or "").lower() if p.style else ""
        text = p.text.strip()
        if not text:
            i += 1
            continue
        if "list" in style_name or style_name.startswith("bullet") or text.startswith(("•", "–", "-", "*")):
            bullets.append(text.lstrip("•–-* ").strip())
        elif not intro:
            intro = text
        else:
            # Additional body paragraph — treat as continued intro or extra bullet
            bullets.append(text)
        i += 1
    return intro, bullets, i


def _collect_table_rows(paragraphs, tables, heading_idx: int) -> tuple[list[list[str]], int]:
    """Find the first table that follows paragraphs[heading_idx].
    Returns (rows_as_lists_of_cell_text, next_heading_index).
    Each row is a list of stripped cell text strings.
    """
    # Find approximate character position of the heading in the body
    # python-docx tables and paragraphs are siblings in body; we locate
    # by walking the raw XML body children.
    from docx.oxml.ns import qn as _qn2  # local import for safety

    body = paragraphs[heading_idx]._element.getparent()
    body_children = list(body)
    heading_el = paragraphs[heading_idx]._element
    try:
        heading_pos = body_children.index(heading_el)
    except ValueError:
        return [], heading_idx + 1

    # Find next <w:tbl> after the heading in body children
    for pos in range(heading_pos + 1, len(body_children)):
        el = body_children[pos]
        if el.tag == _qn2("w:tbl"):
            from docx.table import Table
            tbl = Table(el, paragraphs[heading_idx]._element.getparent())
            rows = []
            for row in tbl.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            # Find next heading index after this table
            next_h = heading_idx + 1
            while next_h < len(paragraphs):
                if _is_heading(paragraphs[next_h]):
                    break
                next_h += 1
            return rows, next_h
        # If we hit another heading paragraph, stop looking
        if el.tag == _qn2("w:p"):
            from docx.text.paragraph import Paragraph
            p = Paragraph(el, body)
            if _is_heading(p):
                break

    return [], heading_idx + 1


# ---------------------------------------------------------------------------
# Doc-type-specific extractors
# ---------------------------------------------------------------------------

def _extract_procedure(paragraphs, tables) -> dict:
    data: dict = {
        "procedure_title": "",
        "doc_number": "",
        "header_date": "",
        "revision": "A",
        "revision_history": [{"rev": "0", "date": "", "description": "Initial issue",
                               "originator": "", "reviewer": "", "approver": ""}],
        "change_log": [],
        "purpose_text": "",
        "scope_text": "",
        "roles_intro": "Roles and responsibilities for this procedure include the following.",
        "roles": [{"role": "All Personnel", "responsibilities": ["Comply with this procedure."]}],
        "ppe_paragraphs": ["Refer to the site-specific PPE matrix for applicable requirements."],
        "procedure_section_title": "",
        "procedure_intro": "",
        "steps_intro": "The responsible party completes the following:",
        "steps": [{"step": "1", "description": "See document body."}],
        "recordkeeping_text": "",
        "terms_intro": "The following terms are specific to this document.",
        "definitions": [],
        "abbreviations_intro": "The following abbreviations and acronyms are specific to this document.",
        "abbreviations": [],
        "references": [],
        "appendices": "",
    }

    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        if not (_is_heading(p)):
            i += 1
            continue
        norm = _normalize_heading(p.text)

        if "PURPOSE" in norm and "SCOPE" not in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["purpose_text"] = text or data["purpose_text"]
        elif "SCOPE" in norm and "PURPOSE" not in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["scope_text"] = text or data["scope_text"]
        elif "PERSONAL PROTECTIVE EQUIPMENT" in norm or norm.endswith("PPE"):
            _, bullets, i = _collect_bullets_until_next_heading(paragraphs, i)
            if bullets:
                data["ppe_paragraphs"] = bullets
        elif "RECORD" in norm and "KEEP" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["recordkeeping_text"] = text or data["recordkeeping_text"]
        elif "DEFINITION" in norm or "TERM" in norm:
            rows, i = _collect_table_rows(paragraphs, tables, i)
            if len(rows) > 1:
                data["definitions"] = [
                    {"term": r[0], "definition": r[1]} for r in rows[1:] if r[0]
                ]
        elif "ABBREVIAT" in norm:
            rows, i = _collect_table_rows(paragraphs, tables, i)
            if len(rows) > 1:
                data["abbreviations"] = [
                    {"abbreviation": r[0], "definition": r[1]} for r in rows[1:] if r[0]
                ]
        elif "REFERENCE" in norm:
            rows, i = _collect_table_rows(paragraphs, tables, i)
            if len(rows) > 1:
                data["references"] = [
                    {"title": r[0], "number": r[1] if len(r) > 1 else ""} for r in rows[1:] if r[0]
                ]
        elif "APPENDIX" in norm or "APPENDICES" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["appendices"] = text
        else:
            # Could be procedure section (5.0 Hot Work Procedure) or roles
            if "ROLE" in norm or "RESPONSIBILIT" in norm:
                intro, _, i = _collect_bullets_until_next_heading(paragraphs, i)
                if intro:
                    data["roles_intro"] = intro
            elif re.match(r"\d+\.\d+\s", p.text.strip()) and not data["procedure_section_title"]:
                # Numbered section that looks like the main procedure section
                data["procedure_section_title"] = p.text.strip()
                intro, _, i = _collect_bullets_until_next_heading(paragraphs, i)
                if intro:
                    data["procedure_intro"] = intro
            else:
                i += 1

    # Use doc heading as title fallback
    for p in paragraphs[:5]:
        if p.style and p.style.name in ("Title", "Heading 1") and p.text.strip():
            if not data["procedure_title"]:
                data["procedure_title"] = p.text.strip()
            break

    return data


def _extract_standard(paragraphs, tables) -> dict:
    data: dict = {
        "document_name": "",
        "doc_number": "",
        "introduction_text": "",
        "scope_text": "",
        "governance_text": "",
        "exception_text": "",
        "continuous_text": "",
        "content_sections": [],
        "definitions": [],
        "references": [],
        "raci": {
            "responsible": "", "accountable": "",
            "consulted": "", "informed": ""
        },
        "approval": {
            "issuer": "", "adopted_by": "NextDecade Corporation", "effective_date": ""
        },
        "revision_history": [{"rev": "0", "description": "Initial issue"}],
    }

    known_headings = set(_STD_HEADING_MAP.keys())
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        if not (_is_heading(p)):
            i += 1
            continue
        norm = _normalize_heading(p.text)

        if "INTRODUCTION" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["introduction_text"] = text
        elif "SCOPE" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["scope_text"] = text
        elif "INTEGRATED GOVERNANCE" in norm or "GOVERNANCE FRAMEWORK" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["governance_text"] = text
        elif "DEFINITION" in norm:
            rows, i = _collect_table_rows(paragraphs, tables, i)
            if len(rows) > 1:
                data["definitions"] = [
                    {"no": str(idx), "term": r[0], "definition": r[1]}
                    for idx, r in enumerate(rows[1:], 1) if r[0]
                ]
        elif "REFERENCE" in norm:
            rows, i = _collect_table_rows(paragraphs, tables, i)
            if len(rows) > 1:
                data["references"] = [
                    {"title": r[0], "number": r[1] if len(r) > 1 else ""}
                    for r in rows[1:] if r[0]
                ]
        elif "EXCEPTION" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["exception_text"] = text
        elif "CONTINUOUS IMPROVEMENT" in norm or "CONTINUOUS" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["continuous_text"] = text
        elif "OWNERSHIP" in norm:
            rows, i = _collect_table_rows(paragraphs, tables, i)
            if rows:
                flat = [c for r in rows for c in r if c]
                if len(flat) >= 4:
                    data["raci"] = {
                        "responsible": flat[0], "accountable": flat[1],
                        "consulted": flat[2], "informed": flat[3],
                    }
        elif "APPROVAL" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["approval"]["issuer"] = text.split("\n")[0] if text else ""
        elif "REVISION HISTORY" in norm:
            rows, i = _collect_table_rows(paragraphs, tables, i)
            if len(rows) > 1:
                data["revision_history"] = [
                    {"rev": r[0], "description": r[1] if len(r) > 1 else ""}
                    for r in rows[1:] if r[0]
                ]
        else:
            # Content section
            intro, bullets, i = _collect_bullets_until_next_heading(paragraphs, i)
            heading_text = p.text.strip()
            if heading_text and any(c.isupper() for c in heading_text):
                # Skip internal structure headings that were missed above
                unmapped = not any(kw in norm for kw in known_headings)
                if unmapped:
                    data["content_sections"].append({
                        "title": heading_text,
                        "intro": intro,
                        "bullets": bullets or ["See document body."],
                    })
                else:
                    pass  # Already handled above — shouldn't reach here

    if not data["content_sections"]:
        data["content_sections"] = [
            {"title": "Requirements", "intro": "", "bullets": ["See document body."]}
        ]

    # Try to find document_name from title paragraph
    for p in paragraphs[:8]:
        if p.style and p.style.name in ("Title", "Heading 1") and p.text.strip():
            data["document_name"] = p.text.strip()
            break

    return data


def _extract_guidance(paragraphs, tables) -> dict:
    data: dict = {
        "document_name": "",
        "doc_number": "",
        "purpose_text": "",
        "governance_text": "",
        "guideline_intro": "",
        "guideline_bullets": ["See document body."],
        "raci": {
            "responsible": "", "accountable": "",
            "consulted": "", "informed": ""
        },
        "approval": {
            "issuer": "", "adopted_by": "NextDecade Corporation", "effective_date": ""
        },
        "revision_history": [{"rev": "0", "description": "Initial issue"}],
    }

    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        if not (_is_heading(p)):
            i += 1
            continue
        norm = _normalize_heading(p.text)

        if "PURPOSE" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["purpose_text"] = text
        elif "INTEGRATED GOVERNANCE" in norm or "GOVERNANCE FRAMEWORK" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["governance_text"] = text
        elif "GUIDELINE" in norm:
            intro, bullets, i = _collect_bullets_until_next_heading(paragraphs, i)
            data["guideline_intro"] = intro
            if bullets:
                data["guideline_bullets"] = bullets
        elif "OWNERSHIP" in norm:
            rows, i = _collect_table_rows(paragraphs, tables, i)
            if rows:
                flat = [c for r in rows for c in r if c]
                if len(flat) >= 4:
                    data["raci"] = {
                        "responsible": flat[0], "accountable": flat[1],
                        "consulted": flat[2], "informed": flat[3],
                    }
        elif "APPROVAL" in norm:
            text, i = _collect_text_until_next_heading(paragraphs, i)
            data["approval"]["issuer"] = text.split("\n")[0] if text else ""
        elif "REVISION HISTORY" in norm:
            rows, i = _collect_table_rows(paragraphs, tables, i)
            if len(rows) > 1:
                data["revision_history"] = [
                    {"rev": r[0], "description": r[1] if len(r) > 1 else ""}
                    for r in rows[1:] if r[0]
                ]
        else:
            i += 1

    for p in paragraphs[:8]:
        if p.style and p.style.name in ("Title", "Heading 1") and p.text.strip():
            data["document_name"] = p.text.strip()
            break

    return data


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def _extract_images(docx_path: Path, images_dir: Path) -> list[dict]:
    """Extract all embedded images from the .docx zip to images_dir.

    Returns a list of dicts: {"filename": "img_001.png", "original": "image1.png",
    "after_heading": "SECTION NAME OR EMPTY"}.

    Determines the preceding heading for each image by walking the document
    body XML and noting which <w:p> with a Heading style appears before the
    paragraph containing the <w:drawing>/<v:imagedata> element.
    """
    images_dir.mkdir(parents=True, exist_ok=True)
    records: list[dict] = []

    with zipfile.ZipFile(docx_path) as z:
        media_files = sorted(n for n in z.namelist() if n.startswith("word/media/"))
        if not media_files:
            return []

        # Build relationship map: rId → media path
        rel_map: dict[str, str] = {}
        for relfile in ("word/_rels/document.xml.rels",):
            if relfile in z.namelist():
                import xml.etree.ElementTree as ET
                root = ET.fromstring(z.read(relfile))
                ns = "http://schemas.openxmlformats.org/package/2006/relationships"
                for rel in root.findall(f"{{{ns}}}Relationship"):
                    target = rel.get("Target", "")
                    rid = rel.get("Id", "")
                    if target.startswith("media/"):
                        rel_map[rid] = "word/" + target

        # Parse document body to find heading → image order
        doc_xml = z.read("word/document.xml").decode("utf-8")

        # Extract all image rIds in document order with their preceding heading
        import xml.etree.ElementTree as ET
        root = ET.fromstring(doc_xml)
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        A_DRAW = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        A_BLI = "http://schemas.openxmlformats.org/drawingml/2006/picture"
        V_NS = "urn:schemas-microsoft-com:vml"
        R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        def _iter_body(elem):
            """Flatten all <w:p> and <w:tbl> at any depth."""
            for child in elem:
                if child.tag == f"{{{W}}}body":
                    yield from _iter_body(child)
                elif child.tag in (f"{{{W}}}p", f"{{{W}}}tbl"):
                    yield child
                else:
                    yield from _iter_body(child)

        current_heading = ""
        image_order: list[tuple[str, str]] = []  # (rId, heading_before)

        for p_el in _iter_body(root):
            if p_el.tag != f"{{{W}}}p":
                continue
            # Check if heading
            pPr = p_el.find(f"{{{W}}}pPr")
            if pPr is not None:
                pStyle = pPr.find(f"{{{W}}}pStyle")
                if pStyle is not None:
                    style_val = pStyle.get(f"{{{W}}}val", "")
                    if "Heading" in style_val or style_val.startswith("heading"):
                        texts = "".join(t.text or "" for t in p_el.iter(f"{{{W}}}t"))
                        if texts.strip():
                            current_heading = _normalize_heading(texts.strip())

            # Check for images in this paragraph
            for drawing in p_el.iter(f"{{{W}}}drawing"):
                for blip in drawing.iter(f"{{{A_BLI}}}blipFill"):
                    for b in blip:
                        rid = b.get(f"{{{R_NS}}}embed", "")
                        if rid:
                            image_order.append((rid, current_heading))
            # Legacy VML images
            for imgdata in p_el.iter(f"{{{V_NS}}}imagedata"):
                rid = imgdata.get(f"{{{R_NS}}}id", "")
                if rid:
                    image_order.append((rid, current_heading))

        # Extract media files in document order
        used_rids: list[str] = []
        for rid, heading in image_order:
            if rid not in rel_map:
                continue
            media_path = rel_map[rid]
            if media_path not in z.namelist():
                continue
            suffix = Path(media_path).suffix.lower() or ".png"
            seq = len(records) + 1
            out_name = f"img_{seq:03d}{suffix}"
            out_path = images_dir / out_name
            with open(out_path, "wb") as f:
                f.write(z.read(media_path))
            records.append({
                "filename": out_name,
                "original": Path(media_path).name,
                "after_heading": heading,
                "path": str(out_path),
            })
            used_rids.append(media_path)

        # Any media files not found via relationships (e.g. EMF thumbs)
        for mf in media_files:
            if mf in used_rids:
                continue
            suffix = Path(mf).suffix.lower() or ".bin"
            seq = len(records) + 1
            out_name = f"img_{seq:03d}{suffix}"
            out_path = images_dir / out_name
            with open(out_path, "wb") as f:
                f.write(z.read(mf))
            records.append({
                "filename": out_name,
                "original": Path(mf).name,
                "after_heading": "",
                "path": str(out_path),
            })

    return records


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract(
    docx_path: str | Path,
    output_json: str | Path | None = None,
    images_dir: str | Path | None = None,
) -> dict:
    """Extract a .docx document into a schema-compatible JSON + images.

    Returns a report dict:
      {
        "doc_type": "procedure" | "standard" | "guidance",
        "output_json": "/path/to/output.json" or None,
        "images": [{"filename": "img_001.png", "after_heading": "SCOPE", ...}],
        "data": { ... }   # the extracted JSON data dict
      }
    """
    if Document is None:
        raise ImportError("python-docx is required: pip install python-docx")

    docx_path = Path(docx_path)
    doc = Document(str(docx_path))

    # Use doc.paragraphs — these are properly linked to the document part so
    # p.style.name works correctly. doc.paragraphs covers direct <w:body>
    # children; SDT-wrapped TOC paragraphs are omitted but that's fine since
    # section headings are never inside SDTs.
    para_objs = doc.paragraphs

    doc_type = _detect_doc_type(para_objs)

    tables = doc.tables
    if doc_type == "procedure":
        data = _extract_procedure(para_objs, tables)
    elif doc_type == "standard":
        data = _extract_standard(para_objs, tables)
    else:
        data = _extract_guidance(para_objs, tables)

    # Extract images
    if images_dir is None and output_json is not None:
        images_dir = Path(output_json).parent / "images"
    elif images_dir is None:
        images_dir = docx_path.parent / "images"
    images_dir = Path(images_dir)

    image_records = _extract_images(docx_path, images_dir)

    # Add images reference to data (optional field — render pipeline uses it)
    if image_records:
        data["images"] = [
            {"path": r["path"], "after_section": r["after_heading"],
             "width_cm": 14, "caption": ""}
            for r in image_records
        ]

    # Write JSON
    if output_json is not None:
        Path(output_json).write_text(json.dumps(data, indent=2, ensure_ascii=False))

    return {
        "doc_type": doc_type,
        "output_json": str(output_json) if output_json else None,
        "images": image_records,
        "data": data,
    }


def main():
    import argparse
    ap = argparse.ArgumentParser(
        description="Extract a NextDecade .docx document to schema-compatible JSON."
    )
    ap.add_argument("docx", help="Input .docx file path")
    ap.add_argument("output_json", nargs="?", help="Output JSON path (default: stdout)")
    ap.add_argument("--images-dir", default=None, metavar="DIR",
                    help="Directory to extract embedded images into")
    args = ap.parse_args()

    result = extract(
        args.docx,
        output_json=args.output_json,
        images_dir=args.images_dir,
    )

    if args.output_json:
        print(json.dumps({
            "doc_type": result["doc_type"],
            "output_json": result["output_json"],
            "image_count": len(result["images"]),
            "images": result["images"],
        }, indent=2))
    else:
        # Stdout mode: print just the data JSON
        print(json.dumps(result["data"], indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
